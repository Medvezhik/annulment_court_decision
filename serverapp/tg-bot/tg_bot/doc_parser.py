import os.path
import logging
from typing import List
import subprocess

import docx
import requests
from docx.opc.exceptions import OpcError

from tg_bot.config import Config
from tg_bot.exceptions import BotException, UnsupportedFileTypeException


logger = logging.getLogger(name=__name__)


def is_tg_file_court_decision(file_id, bot, cfg: Config):
    file_info = bot.get_file(file_id)
    _, ext = os.path.splitext(file_info.file_path)
    if ext.lower() not in [".doc", ".docx"]:
        raise BotException(f"Unsupported file format: {ext}!")

    try:
        response = requests.get(
            f"https://api.telegram.org/file/bot{cfg.tg_api_token}/{file_info.file_path}",
            stream=True,
        )
    except requests.exceptions.RequestException as e:
        logger.error(
            f"Failed to retrieve information about the file at path '{file_info.file_path}': {e}"
        )
        raise BotException("Failed to retrieve information about the file")

    if not os.path.exists(cfg.ms_docs_dir):
        os.makedirs(cfg.ms_docs_dir, exist_ok=True)
    tmp_file = os.path.join(cfg.ms_docs_dir, f"{file_id}{ext}")
    with open(tmp_file, "wb") as f:
        f.write(response.raw.read())

    return is_ms_file_court_decision(tmp_file, cfg.court_required_words)


def get_text_from_tg_file(file_id, bot, cfg: Config) -> str:
    file_info = bot.get_file(file_id)
    _, ext = os.path.splitext(file_info.file_path)
    if ext.lower() not in [".doc", ".docx"]:
        raise UnsupportedFileTypeException()

    try:
        response = requests.get(
            f"https://api.telegram.org/file/bot{cfg.tg_api_token}/{file_info.file_path}",
            stream=True,
        )
    except requests.exceptions.RequestException as e:
        logger.error(
            f"Failed to retrieve information about the file at path '{file_info.file_path}': {e}"
        )
        raise BotException("Failed to retrieve information about the file")

    if not os.path.exists(cfg.ms_docs_dir):
        os.makedirs(cfg.ms_docs_dir, exist_ok=True)
    tmp_file = os.path.join(cfg.ms_docs_dir, f"{file_id}{ext}")
    with open(tmp_file, "wb") as f:
        f.write(response.raw.read())

    return text_from_ms_word_file(path=tmp_file)


def text_from_ms_word_file(path) -> str:
    text = ""
    try:
        ms_doc = docx.Document(path)
        text = " ".join(p.text for p in ms_doc.paragraphs)
    except OpcError as e:
        logger.error(
            f"Failed to parse file {path} with docx lib: {e}. Extracting text with 'antiword'..."
        )
        # docx-python cannot parse some doc files.
        # Attempting to extract text from those files with antiword utility:
        process_res = subprocess.run(["antiword", path], capture_output=True, text=True)
        if process_res.returncode != 0:
            logger.error(f"Failed to parse file with antiword: {process_res.stderr}")
            raise BotException(f"Failed to parse file {path}")
        # Replace all newline characters with spaces:
        text = process_res.stdout.replace("\n", " ")

    logging.debug(f"Successfully parsed the file {path}")
    return text


def is_text_court_decision(text: str, key_words: List[str]):
    for w in key_words:
        if w in text:
            print(f"Detected key word '{w}' in text")
            return True
    return False


def is_ms_file_court_decision(path, key_words: List[str]):
    try:
        ms_doc = docx.Document(path)
    except OpcError as e:
        logger.error(f"Failed to parse file {path} with docx library: {e}")
        raise BotException(f"Failed to parse file {path}")

    for p in ms_doc.paragraphs:
        for w in key_words:
            if w in p.text:
                return True
    return False
